"""
Módulo mejorado para manejar documentos Word (.docx) con detección automática de formatos
"""
import anthropic
import os
import json
import re
from pathlib import Path
from typing import Dict, List, Tuple
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

class WordHandler:
    def __init__(self, api_key: str = None):
        """
        Inicializa el manejador de Word con Claude API

        Args:
            api_key: API key de Anthropic
        """
        self.api_key = api_key or os.getenv('ANTHROPIC_API_KEY')
        if not self.api_key:
            raise ValueError("Se requiere ANTHROPIC_API_KEY")

        self.client = anthropic.Anthropic(api_key=self.api_key)

    def extraer_texto_word(self, docx_path: str) -> str:
        """
        Extrae todo el texto de un documento Word

        Args:
            docx_path: Ruta al archivo .docx

        Returns:
            Texto completo del documento
        """
        doc = Document(docx_path)
        texto_completo = []

        for para in doc.paragraphs:
            texto_completo.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    texto_completo.append(cell.text)

        return '\n'.join(texto_completo)

    def detectar_tipo_campos(self, texto: str) -> str:
        """
        Detecta qué tipo de campos tiene el documento

        Returns:
            Tipo de campos detectado: 'marcadores', 'puntos_suspensivos', 'lineas', 'haga_clic', 'mixto'
        """
        tiene_marcadores = bool(re.search(r'\{\{[A-Z_]+\}\}', texto))
        tiene_puntos = bool(re.search(r'\.{3,}|…{2,}', texto))
        tiene_lineas = bool(re.search(r'_{3,}', texto))
        tiene_haga_clic = 'Haga clic aquí' in texto or 'haga clic' in texto.lower()

        if tiene_marcadores:
            return 'marcadores'
        elif tiene_haga_clic:
            return 'haga_clic'
        elif tiene_puntos:
            return 'puntos_suspensivos'
        elif tiene_lineas:
            return 'lineas'
        else:
            return 'mixto'

    def extraer_datos_cliente_word(self, docx_path: str) -> Dict[str, any]:
        """
        Extrae datos del cliente desde un documento Word usando Claude API

        Args:
            docx_path: Ruta al archivo Word

        Returns:
            Diccionario con los datos extraídos
        """
        texto = self.extraer_texto_word(docx_path)

        prompt = f"""Analiza este texto extraído de un documento Word y extrae la siguiente información sobre el cliente/empresa:

TEXTO DEL DOCUMENTO:
{texto}

CAMPOS A EXTRAER:
1. nombre_representante_legal: Nombre completo del representante legal
2. dni_representante: DNI/NIF del representante legal
3. razon_social: Razón social de la empresa
4. cif: CIF de la empresa
5. direccion: Dirección completa de la empresa
6. correo_electronico: Correo electrónico de contacto
7. numero_trabajadores: Número de trabajadores (como número entero)
8. facturacion: Facturación anual (como número decimal, sin símbolos)
9. habilitaciones: Lista de habilitaciones (separadas por comas)
10. isos: Certificaciones ISO que posee (separadas por comas)
11. rolece: Número o código ROLECE si está presente
12. tiene_plan_igualdad: true si tiene plan de igualdad, false si no
13. tiene_protocolo_acoso: true si tiene protocolo de acoso, false si no

Responde ÚNICAMENTE con un objeto JSON válido. Si un campo no está presente, usa null.

Ejemplo:
{{
  "nombre_representante_legal": "María López",
  "dni_representante": "12345678A",
  "razon_social": "Empresa SA",
  "cif": "A12345678",
  "direccion": "Calle Principal 1",
  "correo_electronico": "info@empresa.com",
  "numero_trabajadores": 25,
  "facturacion": 500000.00,
  "habilitaciones": "Transporte, Logística",
  "isos": "ISO 9001",
  "rolece": null,
  "tiene_plan_igualdad": false,
  "tiene_protocolo_acoso": true
}}"""

        try:
            message = self.client.messages.create(
                model="claude-3-haiku-20240307",
                max_tokens=2048,
                messages=[
                    {
                        "role": "user",
                        "content": prompt
                    }
                ]
            )

            response_text = message.content[0].text

            if "```json" in response_text:
                response_text = response_text.split("```json")[1].split("```")[0]
            elif "```" in response_text:
                response_text = response_text.split("```")[1].split("```")[0]

            datos = json.loads(response_text.strip())
            datos['pdf_original_nombre'] = Path(docx_path).name
            datos['pdf_original_ruta'] = docx_path

            return datos

        except Exception as e:
            raise Exception(f"Error al extraer datos del Word: {e}")

    def analizar_campos_con_ia(self, texto_doc: str, datos_cliente: Dict) -> List[Dict]:
        """
        Usa IA para identificar qué datos del cliente van en qué parte del documento

        Returns:
            Lista de mapeos: [{"texto_buscar": "...", "reemplazar_con": "..."}]
        """
        datos_json = json.dumps(datos_cliente, indent=2, ensure_ascii=False)

        prompt = f"""Analiza este documento y crea un mapeo de qué campos deben rellenarse con qué datos del cliente.

DOCUMENTO:
{texto_doc[:3000]}

DATOS DEL CLIENTE:
{datos_json}

Identifica:
1. Campos vacíos marcados con: … (puntos suspensivos), ____ (líneas), "Haga clic aquí", etc.
2. Etiquetas antes de los campos (ej: "D./Dña", "con DNI número", "Razón social:", etc.)
3. Checkboxes (☐) que deben marcarse según datos booleanos

Responde con JSON en este formato:
{{
  "reemplazos": [
    {{
      "patron": "D./Dña.*?(?=,)",
      "contexto": "D./Dña",
      "valor": "Juan Pérez García",
      "tipo": "texto"
    }},
    {{
      "patron": "☐.*?50 o más trabajadores",
      "contexto": "50 o más trabajadores",
      "valor": "[X]",
      "tipo": "checkbox"
    }}
  ]
}}

IMPORTANTE: Responde SOLO con JSON válido."""

        try:
            message = self.client.messages.create(
                model="claude-3-haiku-20240307",
                max_tokens=3000,
                messages=[{"role": "user", "content": prompt}]
            )

            response_text = message.content[0].text
            if "```json" in response_text:
                response_text = response_text.split("```json")[1].split("```")[0]
            elif "```" in response_text:
                response_text = response_text.split("```")[1].split("```")[0]

            resultado = json.loads(response_text.strip())
            return resultado.get('reemplazos', [])

        except Exception as e:
            print(f"Error en análisis IA: {e}")
            return []

    def rellenar_word_inteligente(self, docx_path: str, datos_cliente: Dict, output_path: str) -> Dict:
        """
        Método inteligente que mantiene el formato original y rellena campos automáticamente

        Args:
            docx_path: Ruta al documento Word original
            datos_cliente: Datos del cliente
            output_path: Ruta de salida

        Returns:
            Información del proceso
        """
        # Cargar documento original
        doc = Document(docx_path)
        texto_completo = self.extraer_texto_word(docx_path)

        # Detectar tipo de campos
        tipo_campos = self.detectar_tipo_campos(texto_completo)

        print(f"Tipo de campos detectado: {tipo_campos}")

        # Crear mapeo inteligente de reemplazos
        reemplazos = self._crear_mapeo_reemplazos(datos_cliente)

        # Aplicar reemplazos en párrafos
        for para in doc.paragraphs:
            texto_original = para.text
            texto_nuevo = texto_original

            for patron, valor in reemplazos.items():
                texto_nuevo = texto_nuevo.replace(patron, str(valor) if valor else '')

            if texto_nuevo != texto_original:
                para.text = texto_nuevo

        # Aplicar reemplazos en tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        texto_original = para.text
                        texto_nuevo = texto_original

                        for patron, valor in reemplazos.items():
                            texto_nuevo = texto_nuevo.replace(patron, str(valor) if valor else '')

                        if texto_nuevo != texto_original:
                            para.text = texto_nuevo

        # Guardar documento
        doc.save(output_path)

        return {
            'exito': True,
            'metodo': f'inteligente_{tipo_campos}',
            'mensaje': f'Documento rellenado (tipo: {tipo_campos})'
        }

    def _crear_mapeo_reemplazos(self, datos_cliente: Dict) -> Dict[str, str]:
        """
        Crea un diccionario de patrones a reemplazar

        Returns:
            Diccionario {patron: valor}
        """
        reemplazos = {}

        # Patrones comunes con marcadores
        reemplazos['{{NOMBRE_REPRESENTANTE}}'] = datos_cliente.get('nombre_representante_legal', '')
        reemplazos['{{DNI_REPRESENTANTE}}'] = datos_cliente.get('dni_representante', '')
        reemplazos['{{RAZON_SOCIAL}}'] = datos_cliente.get('razon_social', '')
        reemplazos['{{CIF}}'] = datos_cliente.get('cif', '')
        reemplazos['{{DIRECCION}}'] = datos_cliente.get('direccion', '')
        reemplazos['{{EMAIL}}'] = datos_cliente.get('correo_electronico', '')
        reemplazos['{{NUM_TRABAJADORES}}'] = str(datos_cliente.get('numero_trabajadores', ''))
        reemplazos['{{FACTURACION}}'] = str(datos_cliente.get('facturacion', ''))

        # Patrones con puntos suspensivos (diferentes variaciones)
        nombre = datos_cliente.get('nombre_representante_legal', '')
        dni = datos_cliente.get('dni_representante', '')
        razon_social = datos_cliente.get('razon_social', '')
        cif = datos_cliente.get('cif', '')
        direccion = datos_cliente.get('direccion', '')

        # Reemplazos para formato "D....……, vecino de.……..."
        if nombre:
            reemplazos['Haga clic aquí para escribir texto'] = nombre
            for i in range(3, 30):
                reemplazos['.' * i] = ''  # Primero limpiamos puntos
                reemplazos['…' * i] = ''  # Y puntos suspensivos unicode

        # Después de limpiar, añadimos valores específicos según contexto
        # Estos se aplicarán de forma inteligente
        self._valores_disponibles = {
            'nombre': nombre,
            'dni': dni,
            'razon_social': razon_social,
            'cif': cif,
            'direccion': direccion,
            'num_trabajadores': datos_cliente.get('numero_trabajadores', ''),
            'plan_igualdad': 'Sí' if datos_cliente.get('tiene_plan_igualdad') else 'No',
            'protocolo_acoso': 'Sí' if datos_cliente.get('tiene_protocolo_acoso') else 'No'
        }

        return reemplazos

    def rellenar_word_con_ia(self, docx_path: str, datos_cliente: Dict, output_path: str) -> Dict:
        """
        Usa IA para analizar el documento completo y rellenarlo manteniendo formato

        Args:
            docx_path: Ruta al documento Word
            datos_cliente: Datos del cliente
            output_path: Ruta de salida

        Returns:
            Información del proceso
        """
        # Extraer texto original
        texto_original = self.extraer_texto_word(docx_path)
        datos_json = json.dumps(datos_cliente, indent=2, ensure_ascii=False)

        # Análisis con IA para crear documento rellenado
        prompt = f"""Tienes un documento Word y datos de un cliente. Rellena el documento de forma inteligente.

DOCUMENTO ORIGINAL:
{texto_original}

DATOS DEL CLIENTE:
{datos_json}

INSTRUCCIONES:
1. Identifica campos vacíos: puntos suspensivos (…), líneas (_____), "Haga clic aquí", etc.
2. Rellena cada campo con el dato correspondiente del cliente
3. Para checkboxes (☐), ponles [X] si el dato es true
4. Mantén TODA la estructura y formato del documento
5. Devuelve el documento COMPLETO rellenado

Responde SOLO con el texto del documento rellenado, sin explicaciones."""

        try:
            message = self.client.messages.create(
                model="claude-3-haiku-20240307",
                max_tokens=4096,
                messages=[{"role": "user", "content": prompt}]
            )

            texto_rellenado = message.content[0].text

            # Crear documento preservando estructura original
            doc_original = Document(docx_path)
            lineas_nuevas = texto_rellenado.split('\n')
            idx_linea = 0

            # Reemplazar párrafos manteniendo formato
            for para in doc_original.paragraphs:
                if idx_linea < len(lineas_nuevas):
                    # Mantener formato del párrafo original
                    for run in para.runs:
                        run.text = ''
                    if para.runs:
                        para.runs[0].text = lineas_nuevas[idx_linea]
                    else:
                        para.add_run(lineas_nuevas[idx_linea])
                    idx_linea += 1

            doc_original.save(output_path)

            return {
                'exito': True,
                'mensaje': 'Documento Word rellenado con IA'
            }

        except Exception as e:
            raise Exception(f"Error al rellenar Word con IA: {e}")

    def rellenar_word(self, docx_path: str, datos_cliente: Dict, output_path: str) -> Dict:
        """
        Método principal mejorado que usa IA para análisis robusto

        Args:
            docx_path: Ruta al documento Word
            datos_cliente: Datos del cliente
            output_path: Ruta de salida

        Returns:
            Información del proceso
        """
        # Usar método con IA directamente (más robusto con cualquier formato)
        return self.rellenar_word_con_ia(docx_path, datos_cliente, output_path)

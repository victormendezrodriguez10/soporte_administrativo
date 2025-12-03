"""
Script para crear una plantilla de ejemplo de Word con marcadores
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def crear_plantilla_word():
    """Crea un documento Word de ejemplo con todos los marcadores"""

    doc = Document()

    # Título
    titulo = doc.add_heading('FORMULARIO DE DATOS DEL CLIENTE', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Sección 1: Datos de la Empresa
    doc.add_heading('1. DATOS DE LA EMPRESA', level=1)

    doc.add_paragraph(f'Razón Social: {{{{RAZON_SOCIAL}}}}')
    doc.add_paragraph(f'CIF: {{{{CIF}}}}')
    doc.add_paragraph(f'Dirección: {{{{DIRECCION}}}}')
    doc.add_paragraph(f'Correo Electrónico: {{{{EMAIL}}}}')

    doc.add_paragraph()

    # Sección 2: Representante Legal
    doc.add_heading('2. REPRESENTANTE LEGAL', level=1)

    doc.add_paragraph(f'Nombre Completo: {{{{NOMBRE_REPRESENTANTE}}}}')
    doc.add_paragraph(f'DNI/NIF: {{{{DNI_REPRESENTANTE}}}}')

    doc.add_paragraph()

    # Sección 3: Datos Operacionales
    doc.add_heading('3. DATOS OPERACIONALES', level=1)

    doc.add_paragraph(f'Número de Trabajadores: {{{{NUM_TRABAJADORES}}}}')
    doc.add_paragraph(f'Facturación Anual: {{{{FACTURACION}}}} €')

    doc.add_paragraph()

    # Sección 4: Certificaciones y Habilitaciones
    doc.add_heading('4. CERTIFICACIONES Y HABILITACIONES', level=1)

    doc.add_paragraph(f'Habilitaciones: {{{{HABILITACIONES}}}}')
    doc.add_paragraph(f'Certificaciones ISO: {{{{ISOS}}}}')
    doc.add_paragraph(f'Número ROLECE: {{{{ROLECE}}}}')

    doc.add_paragraph()

    # Sección 5: Políticas y Protocolos
    doc.add_heading('5. POLÍTICAS Y PROTOCOLOS', level=1)

    doc.add_paragraph(f'¿Dispone de Plan de Igualdad?: {{{{PLAN_IGUALDAD}}}}')
    doc.add_paragraph(f'¿Dispone de Protocolo de Acoso?: {{{{PROTOCOLO_ACOSO}}}}')

    doc.add_paragraph()
    doc.add_paragraph()

    # Pie de página
    doc.add_paragraph('_' * 50)
    footer = doc.add_paragraph('Documento generado automáticamente por Sistema de Soporte Administrativo')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_format = footer.runs[0].font
    footer_format.size = Pt(9)
    footer_format.italic = True

    # Guardar documento
    output_path = 'plantilla_ejemplo_marcadores.docx'
    doc.save(output_path)
    print(f"✅ Plantilla creada: {output_path}")

    return output_path


def crear_plantilla_tabla():
    """Crea una plantilla con formato de tabla"""

    doc = Document()

    # Título
    titulo = doc.add_heading('FICHA DE CLIENTE', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Crear tabla
    table = doc.add_table(rows=13, cols=2)
    table.style = 'Light Grid Accent 1'

    # Llenar tabla con marcadores
    datos = [
        ('Razón Social', '{{RAZON_SOCIAL}}'),
        ('CIF', '{{CIF}}'),
        ('Dirección', '{{DIRECCION}}'),
        ('Email', '{{EMAIL}}'),
        ('Representante Legal', '{{NOMBRE_REPRESENTANTE}}'),
        ('DNI Representante', '{{DNI_REPRESENTANTE}}'),
        ('Nº Trabajadores', '{{NUM_TRABAJADORES}}'),
        ('Facturación', '{{FACTURACION}}'),
        ('Habilitaciones', '{{HABILITACIONES}}'),
        ('ISOs', '{{ISOS}}'),
        ('ROLECE', '{{ROLECE}}'),
        ('Plan de Igualdad', '{{PLAN_IGUALDAD}}'),
        ('Protocolo de Acoso', '{{PROTOCOLO_ACOSO}}')
    ]

    for i, (campo, marcador) in enumerate(datos):
        row_cells = table.rows[i].cells
        row_cells[0].text = campo
        row_cells[1].text = marcador

    # Guardar
    output_path = 'plantilla_ejemplo_tabla.docx'
    doc.save(output_path)
    print(f"✅ Plantilla con tabla creada: {output_path}")

    return output_path


if __name__ == "__main__":
    print("🔨 Creando plantillas de ejemplo...")
    print()

    crear_plantilla_word()
    crear_plantilla_tabla()

    print()
    print("✅ Plantillas creadas exitosamente")
    print()
    print("Puedes usar estas plantillas para probar el sistema:")
    print("1. plantilla_ejemplo_marcadores.docx - Formato de documento tradicional")
    print("2. plantilla_ejemplo_tabla.docx - Formato de tabla")
    print()
    print("Los marcadores se rellenarán automáticamente con los datos del cliente.")
